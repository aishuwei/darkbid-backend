"""
backend/app.py - Flask后端主入口
放在 D:\Darkbid\backend\app.py

运行方式：
  cd D:\Darkbid && python backend\app.py

访问测试：
  浏览器打开 http://localhost:5000/api/health
"""
import sys
import os

PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import uuid
import time
import json
from datetime import datetime

from backend.config import (
    UPLOAD_DIR, OUTPUT_DIR, MAX_FILE_SIZE, MAX_PAGES,
    VALID_CODES
)

from core.format_checker import FormatChecker
from utils.annotator import DocumentAnnotator
from utils.report_gen import ReportGenerator

app = Flask(__name__)
CORS(app)

tasks = {}


def generate_task_id():
    return 'task_' + str(int(time.time())) + '_' + uuid.uuid4().hex[:6]


def get_task_dir(task_id):
    task_dir = os.path.join(UPLOAD_DIR, task_id)
    os.makedirs(task_dir, exist_ok=True)
    return task_dir


def get_output_dir(task_id):
    out_dir = os.path.join(OUTPUT_DIR, task_id)
    os.makedirs(out_dir, exist_ok=True)
    return out_dir


def validate_document(file_path):
    size = os.path.getsize(file_path)
    if size > MAX_FILE_SIZE:
        return False, f"文件大小{size/1024/1024:.2f}MB，超过限制1MB"
    try:
        from docx import Document
        doc = Document(file_path)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        estimated_pages = para_count // 35 + 1
        if estimated_pages > MAX_PAGES:
            return False, f"估算页数约{estimated_pages}页，超过限制{MAX_PAGES}页"
    except Exception as e:
        return False, f"文件解析失败: {str(e)}"
    return True, "校验通过"


# ========== API接口 ==========

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok', 'message': '追标猎手后端服务运行中'})


@app.route('/api/verify', methods=['POST'])
def verify_code():
    data = request.get_json() or {}
    code = data.get('code', '').strip().upper()
    if not code:
        return jsonify({'success': False, 'message': '请输入激活码'}), 400
    if code in VALID_CODES:
        return jsonify({'success': True, 'message': '验证成功'})
    return jsonify({'success': False, 'message': '激活码无效'}), 401


@app.route('/api/upload-req-text', methods=['POST'])
def upload_requirement_text():
    data = request.get_json() or {}
    text = data.get('text', '').strip()
    if not text:
        return jsonify({'success': False, 'message': '文本不能为空'}), 400
    if len(text) > 5000:
        return jsonify({'success': False, 'message': '文本超过5000字限制'}), 400

    task_id = generate_task_id()
    task_dir = get_task_dir(task_id)
    req_path = os.path.join(task_dir, 'requirement.txt')
    with open(req_path, 'w', encoding='utf-8') as f:
        f.write(text)

    tasks[task_id] = {
        'status': 'req_uploaded',
        'requirement_type': 'text',
        'requirement_path': req_path,
        'created_at': datetime.now().isoformat()
    }
    return jsonify({'success': True, 'task_id': task_id, 'message': '格式要求已接收'})


@app.route('/api/upload-req-file', methods=['POST'])
def upload_requirement_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有上传文件'}), 400
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '文件名为空'}), 400

    task_id = generate_task_id()
    task_dir = get_task_dir(task_id)
    ext = os.path.splitext(file.filename)[1].lower()
    req_path = os.path.join(task_dir, f'requirement{ext}')
    file.save(req_path)

    tasks[task_id] = {
        'status': 'req_uploaded',
        'requirement_type': 'file',
        'requirement_path': req_path,
        'file_name': file.filename,
        'created_at': datetime.now().isoformat()
    }
    return jsonify({'success': True, 'task_id': task_id, 'message': '文件已接收'})


@app.route('/api/generate-rules', methods=['POST'])
def generate_rules():
    data = request.get_json() or {}
    # 兼容驼峰和下划线
    task_id = data.get('taskId') or data.get('task_id')

    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404

    task = tasks[task_id]
    req_path = task.get('requirement_path')
    if not req_path or not os.path.exists(req_path):
        return jsonify({'success': False, 'message': '格式要求文件不存在'}), 400

    # 读取要求文本
    try:
        if req_path.endswith('.txt'):
            with open(req_path, 'r', encoding='utf-8') as f:
                requirement_text = f.read()
        elif req_path.endswith(('.docx', '.doc')):
            from docx import Document
            doc = Document(req_path)
            requirement_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        else:
            requirement_text = ""
    except Exception as e:
        return jsonify({'success': False, 'message': f'读取文件失败: {str(e)}'}), 500

    # 默认配置
        # 贵州省公共资源交易中心暗标规则（内置，黔发改法规〔2026〕196号）
    default_rules = {
        "document_info": {
            "name": None,
            "description": "贵州省工程建设招标投标暗标评审格式要求（黔发改法规〔2026〕196号）",
            "generated_date": "2026-07-16"
        },
        "page_check": {
            "enabled": True,
            "description": "统一采用A4纸张，文字、图表颜色均设置为黑色，总页数不得超过800页",
            "paper_size": "A4",
            "max_pages": 800,
            "no_blank_pages": None
        },
        "margin_check": {
            "enabled": True,
            "description": "上边距2.0厘米，其余（下、左、右）均为2.5厘米",
            "top_cm": 2.0,
            "bottom_cm": 2.5,
            "left_cm": 2.5,
            "right_cm": 2.5,
            "tolerance_cm": 0.1
        },
        "font_check": {
            "enabled": True,
            "description": "标题及正文部分所有文字均采用宋体四号常规字体，黑色，禁止加粗/加色/倾斜/下划线",
            "chinese_font": "宋体",
            "size_pt": 14,
            "color_hex": "000000",
            "allow_bold": False,
            "allow_italic": False,
            "allow_underline": False,
            "allow_color_change": False
        },
        "paragraph_check": {
            "enabled": True,
            "description": "左对齐，首行缩进2字符，固定值25磅行距，段前段后间距为0，不得有空格",
            "alignment": "left",
            "line_spacing_rule": "exact",
            "line_spacing_pt": 25,
            "line_spacing_tolerance_pt": 0.5,
            "first_line_indent_chars": 2,
            "first_line_indent_tolerance_chars": 0.3,
            "space_before_pt": 0,
            "space_after_pt": 0,
            "space_tolerance_pt": 0.5,
            "allow_space_as_indent": False,
            "no_empty_paragraphs": True,
            "allow_left_indent": None,
            "allow_right_indent": None
        },
        "structure_check": {
            "enabled": True,
            "description": "封面后即为正文，不得设置目录、内封面及空白页；不编页码，不设页眉、页脚",
            "allow_toc": False,
            "allow_header": False,
            "allow_footer": False,
            "allow_page_number": False,
            "allow_cover": False
        },
        "table_check": {
            "enabled": True,
            "description": "图表内文字采用宋体常规五号字体，黑色，禁止加粗/加色/倾斜/下划线",
            "table_alignment": None,
            "table_text_alignment": None,
            "table_indent_none": None,
            "table_spacing_none": None,
            "table_font_family": "宋体",
            "table_font_size_pt": 10.5,
            "table_font_color_hex": "000000",
            "table_allow_bold": False,
            "table_allow_italic": False,
            "table_allow_underline": False,
            "table_allow_color_change": False
        },
        "heading_check": {
            "enabled": True,
            "description": "正文标题序号按阿拉伯数字分级编排：一级为1、2、3……；二级为1.1、1.2……；三级为1.1.1、1.1.2……",
            "identification": {
                "method": "regex",
                "patterns": {
                    "level_1": "^\\d+[、\\.]?",
                    "level_2": "^\\d+\\.\\d+[、\\.]?",
                    "level_3": "^\\d+\\.\\d+\\.\\d+[、\\.]?"
                },
                "style_mapping": None
            },
            "level_rules": {
                "level_1": {
                    "enabled": True,
                    "chinese_font": "宋体",
                    "size_pt": 14,
                    "color_hex": "000000",
                    "allow_bold": False,
                    "allow_italic": False,
                    "allow_underline": False,
                    "allow_color_change": False,
                    "alignment": "left",
                    "number_format": None
                },
                "level_2": {
                    "enabled": True,
                    "chinese_font": "宋体",
                    "size_pt": 14,
                    "color_hex": "000000",
                    "allow_bold": False,
                    "allow_italic": False,
                    "allow_underline": False,
                    "allow_color_change": False,
                    "alignment": "left",
                    "number_format": None
                },
                "level_3": {
                    "enabled": True,
                    "chinese_font": "宋体",
                    "size_pt": 14,
                    "color_hex": "000000",
                    "allow_bold": False,
                    "allow_italic": False,
                    "allow_underline": False,
                    "allow_color_change": False,
                    "alignment": "left",
                    "number_format": None
                },
                "level_4": {
                    "enabled": False,
                    "chinese_font": None,
                    "size_pt": None,
                    "color_hex": None,
                    "allow_bold": None,
                    "allow_italic": None,
                    "allow_underline": None,
                    "allow_color_change": None,
                    "alignment": None,
                    "number_format": None
                },
                "level_5": {
                    "enabled": False,
                    "chinese_font": None,
                    "size_pt": None,
                    "color_hex": None,
                    "allow_bold": None,
                    "allow_italic": None,
                    "allow_underline": None,
                    "allow_color_change": None,
                    "alignment": None,
                    "number_format": None
                },
                "level_6": {
                    "enabled": False,
                    "chinese_font": None,
                    "size_pt": None,
                    "color_hex": None,
                    "allow_bold": None,
                    "allow_italic": None,
                    "allow_underline": None,
                    "allow_color_change": None,
                    "alignment": None,
                    "number_format": None
                },
                "level_7": {
                    "enabled": False,
                    "chinese_font": None,
                    "size_pt": None,
                    "color_hex": None,
                    "allow_bold": None,
                    "allow_italic": None,
                    "allow_underline": None,
                    "allow_color_change": None,
                    "alignment": None,
                    "number_format": None
                }
            },
            "validation": {
                "check_continuity": False,
                "check_hierarchy": False,
                "check_orphan_levels": False
            }
        },
        "punctuation_check": {
            "enabled": True,
            "description": "默认中文标点，中文字符间无空格",
            "require_chinese": True,
            "no_space_between_chars": True
        }
    }

    task_dir = get_task_dir(task_id)
    rules_path = os.path.join(task_dir, 'rules.json')
    with open(rules_path, 'w', encoding='utf-8') as f:
        json.dump(default_rules, f, ensure_ascii=False, indent=2)

    tasks[task_id]['status'] = 'rules_generated'
    tasks[task_id]['rules_path'] = rules_path

    return jsonify({'success': True, 'message': '配置已生成', 'task_id': task_id})


@app.route('/api/upload-doc', methods=['POST'])
def upload_document():
    # 兼容 formData 驼峰和下划线
    task_id = request.form.get('taskId') or request.form.get('task_id')

    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404

    if 'file' not in request.files:
        return jsonify({'success': False, 'message': '没有上传文件'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'message': '文件名为空'}), 400

    if not file.filename.lower().endswith('.docx'):
        return jsonify({'success': False, 'message': '仅支持 .docx 格式'}), 400

    file.seek(0, os.SEEK_END)
    file_size = file.tell()
    file.seek(0)

    if file_size > MAX_FILE_SIZE:
        return jsonify({'success': False, 'message': f'文件大小{file_size/1024/1024:.2f}MB，超过限制1MB'}), 400

    task_dir = get_task_dir(task_id)
    doc_path = os.path.join(task_dir, 'document.docx')
    file.save(doc_path)

    valid, msg = validate_document(doc_path)
    if not valid:
        os.remove(doc_path)
        return jsonify({'success': False, 'message': msg}), 400

    tasks[task_id]['status'] = 'doc_uploaded'
    tasks[task_id]['doc_path'] = doc_path
    tasks[task_id]['doc_name'] = file.filename

    return jsonify({'success': True, 'message': '文件上传成功', 'task_id': task_id})


@app.route('/api/start-check', methods=['POST'])
def start_check():
    data = request.get_json() or {}
    task_id = data.get('taskId') or data.get('task_id')

    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404

    task = tasks[task_id]
    if task.get('status') != 'doc_uploaded':
        return jsonify({'success': False, 'message': '请先上传技术文件'}), 400

    doc_path = task.get('doc_path')
    rules_path = task.get('rules_path')

    if not os.path.exists(doc_path):
        return jsonify({'success': False, 'message': '技术文件不存在'}), 400

    if not rules_path or not os.path.exists(rules_path):
        rules_path = os.path.join(PROJECT_ROOT, 'clients', 'config', 'rules.json')
        if not os.path.exists(rules_path):
            return jsonify({'success': False, 'message': '找不到规则配置文件'}), 500

    tasks[task_id]['status'] = 'checking'

    try:
        checker = FormatChecker(rules_path)
        format_issues = checker.check_document(doc_path)

        annotator = DocumentAnnotator()
        out_dir = get_output_dir(task_id)

        annotated_path = annotator.generate_annotated_copy(
            original_path=doc_path,
            format_issues=format_issues,
            output_dir=out_dir,
            suffix='_批注版'
        )

        report_gen = ReportGenerator()
        report_path = report_gen.generate_html_report(
            file_path=doc_path,
            format_issues=format_issues,
            output_dir=out_dir,
            suffix='_检查报告'
        )

        tasks[task_id]['status'] = 'completed'
        tasks[task_id]['format_issues'] = format_issues
        tasks[task_id]['issue_count'] = len(format_issues)
        tasks[task_id]['annotated_path'] = annotated_path
        tasks[task_id]['report_path'] = report_path
        tasks[task_id]['completed_at'] = datetime.now().isoformat()

        return jsonify({
            'success': True,
            'message': '检查完成',
            'task_id': task_id,
            'issue_count': len(format_issues),
            'passed': len(format_issues) == 0
        })

    except Exception as e:
        import traceback
        tasks[task_id]['status'] = 'failed'
        tasks[task_id]['error'] = str(e)
        tasks[task_id]['traceback'] = traceback.format_exc()
        return jsonify({'success': False, 'message': f'检查失败: {str(e)}'}), 500


@app.route('/api/status/<task_id>', methods=['GET'])
def get_status(task_id):
    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    task = tasks[task_id]
    return jsonify({
        'success': True,
        'task_id': task_id,
        'status': task.get('status'),
        'issue_count': task.get('issue_count', 0)
    })


@app.route('/api/download/<task_id>', methods=['GET'])
def download_result(task_id):
    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404
    task = tasks[task_id]
    if task.get('status') != 'completed':
        return jsonify({'success': False, 'message': '检查尚未完成'}), 400

    return jsonify({
        'success': True,
        'task_id': task_id,
        'files': {
            'report': {'name': os.path.basename(task.get('report_path', ''))},
            'annotated': {'name': os.path.basename(task.get('annotated_path', ''))}
        }
    })


# ========== 【新增】文件流下载接口 ==========
@app.route('/api/file/<task_id>', methods=['GET'])
def download_file(task_id):
    """
    下载检查生成的文件流（HTML报告 / 批注版Word）
    小程序通过 wx.downloadFile 调用此接口获取临时文件路径
    """
    file_type = request.args.get('type', 'report')  # 'report' 或 'annotated'

    if not task_id or task_id not in tasks:
        return jsonify({'success': False, 'message': '任务不存在'}), 404

    task = tasks[task_id]
    if task.get('status') != 'completed':
        return jsonify({'success': False, 'message': '检查尚未完成'}), 400

    if file_type == 'report':
        file_path = task.get('report_path')
        mime_type = 'text/html; charset=utf-8'
        download_name = os.path.basename(file_path) if file_path else '格式检查报告.html'
    elif file_type == 'annotated':
        file_path = task.get('annotated_path')
        mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        download_name = os.path.basename(file_path) if file_path else '批注版.docx'
    else:
        return jsonify({'success': False, 'message': '未知文件类型，仅支持 report / annotated'}), 400

    if not file_path or not os.path.exists(file_path):
        return jsonify({'success': False, 'message': '文件不存在或已被清理'}), 404

    try:
        return send_file(
            file_path,
            mimetype=mime_type,
            as_attachment=False,        # False 便于小程序直接预览
            download_name=download_name
        )
    except Exception as e:
        return jsonify({'success': False, 'message': f'文件读取失败: {str(e)}'}), 500


if __name__ == '__main__':
    print("=" * 50)
    print("  追标猎手 - 暗标格式检查后端服务")
    print("=" * 50)
    print(f"  项目根目录: {PROJECT_ROOT}")
    print(f"  上传目录: {UPLOAD_DIR}")
    print(f"  输出目录: {OUTPUT_DIR}")
    print("=" * 50)
    print("  启动地址: http://0.0.0.0:5000")
    print("  测试地址: http://localhost:5000/api/health")
    print("=" * 50)
    app.run(host='0.0.0.0', port=5000, debug=False)