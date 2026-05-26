import json
import re
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def log(msg):
    """内置简单的日志打印函数，适配 main.py 的日志风格"""
    print(f"[INFO] {msg}")


class FormatChecker:
    """
    Word 暗标格式检查器
    核心改进：
    1. 段落格式（缩进、行距、对齐、段间距）直接读取 XML 原始值，避免 python-docx
       Length/EMU/Twips 单位转换导致的误判。
    2. 对继承值（XML 属性缺失）采取宽容策略，不报错。
    3. 错误提示统一使用中文 Word 常用单位（磅、字符、厘米）。
    4. 增加标点符号检查和字符间空格检查。
    5. 表格内文字同样检查字体、颜色、加粗、倾斜。
    6. 增加空段落检查和左右缩进检查。
    """

    def __init__(self, config_path="config/rules.json"):
        with open(config_path, "r", encoding="utf-8") as f:
            self.config = json.load(f)
        self.errors = []

    # ------------------------------------------------------------------
    # 入口
    # ------------------------------------------------------------------
    def check_document(self, file_path):
        """
        入口函数：检查整个文档，并返回错误列表供 main.py 汇总
        """
        log(f"正在加载文档: {file_path}")
        try:
            doc = Document(file_path)
        except Exception as e:
            log(f"无法打开文档: {e}")
            self.errors.append(f"无法打开文档: {e}")
            return self.errors

        self._check_page_setup(doc)
        self._check_margins(doc)
        self._check_fonts_and_styles(doc)
        self._check_paragraph_format(doc)
        self._check_objects(doc)
        self._check_structure(doc)
        self._check_punctuation(doc)

        return self.errors

    # ------------------------------------------------------------------
    # 辅助：安全读取 XML 属性
    # ------------------------------------------------------------------
    @staticmethod
    def _xml_attr(element, attr_local_name):
        """
        从 XML 元素读取指定属性的原始字符串值。
        命名空间前缀固定为 w (wordprocessingml)。
        属性不存在时返回 None。
        """
        if element is None:
            return None
        return element.get(qn(f"w:{attr_local_name}"))

    @staticmethod
    def _xml_attr_twips(element, attr_local_name):
        """
        读取 XML 属性并转为整数 twips。属性不存在或无法解析时返回 None。
        Word XML 中 w:firstLine / w:left / w:right / w:line / w:before / w:after
        等距离属性直接以 twips 为单元存储。
        """
        val = FormatChecker._xml_attr(element, attr_local_name)
        if val is None:
            return None
        try:
            return int(val)
        except (ValueError, TypeError):
            return None

    # ------------------------------------------------------------------
    # 辅助：检查 run 的字体样式（用于正文和表格）
    # ------------------------------------------------------------------
    def _check_run_font(self, run, para_idx, context="段落"):
        """
        检查单个 run 的字体、字号、颜色、加粗、倾斜、下划线
        :param run: docx run 对象
        :param para_idx: 段落序号（用于报错）
        :param context: 上下文描述（"段落"或"表格"）
        """
        cfg = self.config.get("font_check", {})
        if not cfg.get("enabled"):
            return

        target_font = cfg.get("chinese_font", cfg.get("target_font", "宋体"))
        target_size_pt = cfg.get("size_pt", cfg.get("target_size_pt", 14))
        target_color = cfg.get("color_hex", cfg.get("target_color_hex", "000000"))
        allow_bold = cfg.get("allow_bold", False)
        allow_italic = cfg.get("allow_italic", False)
        allow_underline = cfg.get("allow_underline", False)
        allow_color_change = cfg.get("allow_color_change", False)

        if not run.text.strip():
            return

        # --- 中文字体 ---
        rpr = run._element.rPr
        if rpr is not None:
            rfonts = rpr.rFonts
            if rfonts is not None:
                east_asia = rfonts.get(qn("w:eastAsia"))
                if east_asia and target_font not in east_asia:
                    self.errors.append(
                        f"{context} {para_idx}: 中文字体错误 '{run.text[:6]}...' (当前: {east_asia}, 应为: {target_font})"
                    )
            else:
                if run.font.name and target_font not in run.font.name:
                    self.errors.append(
                        f"{context} {para_idx}: 字体错误 '{run.text[:6]}...' (当前: {run.font.name}, 应为: {target_font})"
                    )

        # --- 字号 ---
        if run.font.size is not None:
            actual_pt = run.font.size.pt
            if abs(actual_pt - target_size_pt) > 0.5:
                self.errors.append(
                    f"{context} {para_idx}: 字号错误 (当前: {actual_pt} pt, 应为: {target_size_pt} pt)"
                )

        # --- 颜色 ---
        if not allow_color_change:
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb)
                if color_hex.upper() != target_color.upper():
                    self.errors.append(
                        f"{context} {para_idx}: 字体颜色错误 (当前: {color_hex}, 必须为黑色 {target_color})"
                    )

        # --- 加粗 / 倾斜 / 下划线 ---
        if not allow_bold and run.bold:
            self.errors.append(
                f"{context} {para_idx}: 发现违规加粗内容 -> '{run.text.strip()[:10]}'"
            )
        if not allow_italic and run.italic:
            self.errors.append(
                f"{context} {para_idx}: 发现违规倾斜内容 -> '{run.text.strip()[:10]}'"
            )
        if not allow_underline and run.underline:
            self.errors.append(
                f"{context} {para_idx}: 发现违规下划线内容 -> '{run.text.strip()[:10]}'"
            )

    # ------------------------------------------------------------------
    # 辅助：检查字符间空格
    # ------------------------------------------------------------------
    def _check_spaces_in_text(self, text, para_idx, context="段落"):
        """
        检查文本中中文字符之间是否有空格
        规则：中文字符之间不得出现空格，但英文单词内部、数字与单位之间允许
        """
        cfg = self.config.get("paragraph_check", {})
        if not cfg.get("no_space_between_chars", False):
            return

        if not text or not text.strip():
            return

        # 正则：匹配中文字符 + 空格 + 中文字符
        pattern = re.compile(r'([\u4e00-\u9fa5])[ \u00a0]([\u4e00-\u9fa5])')
        matches = list(pattern.finditer(text))

        if matches:
            # 收集所有违规位置
            positions = []
            for m in matches:
                start = max(0, m.start() - 2)
                end = min(len(text), m.end() + 2)
                snippet = text[start:end]
                positions.append(snippet)

            unique_snippets = list(dict.fromkeys(positions))[:3]  # 去重，最多显示3处
            snippets_str = "; ".join([f"'{s}'" for s in unique_snippets])
            self.errors.append(
                f"{context} {para_idx}: 发现中文字符间有空格 -> {snippets_str}，请删除"
            )

    # ------------------------------------------------------------------
    # 0. 纸张大小与页数
    # ------------------------------------------------------------------
    def _check_page_setup(self, doc):
        cfg = self.config.get("page_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查纸张大小和页数...")
        max_pages = cfg.get("max_pages", 200)
        para_count = len([p for p in doc.paragraphs if p.text.strip()])
        estimated_pages = para_count // 35 + 1
        if estimated_pages > max_pages:
            self.errors.append(
                f"文档页数可能超标 (估算约 {estimated_pages} 页，限制 {max_pages} 页)，请手动核实。"
            )

        for i, section in enumerate(doc.sections):
            w_cm = section.page_width.cm
            h_cm = section.page_height.cm
            is_a4 = (
                (abs(w_cm - 21.0) < 0.5 and abs(h_cm - 29.7) < 0.5)
                or (abs(w_cm - 29.7) < 0.5 and abs(h_cm - 21.0) < 0.5)
            )
            if not is_a4:
                self.errors.append(
                    f"第 {i + 1} 节：纸张大小错误 (当前: {w_cm:.1f}×{h_cm:.1f} cm)，必须为 A4"
                )

    # ------------------------------------------------------------------
    # 1. 页边距
    # ------------------------------------------------------------------
    def _check_margins(self, doc):
        cfg = self.config.get("margin_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查页边距...")
        targets = {
            "top": cfg.get("top_cm", 0),
            "bottom": cfg.get("bottom_cm", 0),
            "left": cfg.get("left_cm", 0),
            "right": cfg.get("right_cm", 0),
        }
        tolerance = cfg.get("tolerance_cm", 0.1)

        for i, section in enumerate(doc.sections):
            actuals = {
                "top": section.top_margin.cm,
                "bottom": section.bottom_margin.cm,
                "left": section.left_margin.cm,
                "right": section.right_margin.cm,
            }
            for name, target in targets.items():
                actual = actuals[name]
                if abs(actual - target) > tolerance:
                    self.errors.append(
                        f"第 {i + 1} 节：{name}边距错误 (当前: {actual:.2f} cm, 要求: {target} cm)"
                    )

    # ------------------------------------------------------------------
    # 2. 字体、颜色、样式（正文 + 表格）
    # ------------------------------------------------------------------
    def _check_fonts_and_styles(self, doc):
        cfg = self.config.get("font_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查字体、字号及样式...")

        # 检查正文段落
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            for run in para.runs:
                self._check_run_font(run, i + 1, "段落")
            # 检查字符间空格
            self._check_spaces_in_text(para.text, i + 1, "段落")

        # 检查表格内文字（如果配置启用）
        if cfg.get("check_table_text", False):
            log("正在检查表格内文字样式...")
            for ti, table in enumerate(doc.tables):
                for row in table.rows:
                    for cell in row.cells:
                        for pi, para in enumerate(cell.paragraphs):
                            if not para.text.strip():
                                continue
                            for run in para.runs:
                                self._check_run_font(run, ti + 1, "表格")
                            # 检查表格内字符间空格
                            self._check_spaces_in_text(para.text, ti + 1, "表格")

    # ------------------------------------------------------------------
    # 3. 段落格式（核心改进区）
    # ------------------------------------------------------------------
    def _check_paragraph_format(self, doc):
        cfg = self.config.get("paragraph_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查段落格式...")

        # 目标值
        target_align = cfg.get("alignment", "left")
        target_line_spacing_pt = cfg.get("line_spacing_pt", cfg.get("line_spacing_value_pt", 30))
        line_spacing_tol_pt = cfg.get("line_spacing_tolerance_pt", 0.5)
        target_indent_chars = cfg.get("first_line_indent_chars", 2)
        indent_tol_chars = cfg.get("first_line_indent_tolerance_chars", 0.3)
        space_before_pt = cfg.get("space_before_pt", 0)
        space_after_pt = cfg.get("space_after_pt", 0)
        space_tol_pt = cfg.get("space_tolerance_pt", 0.5)
        allow_space_as_indent = cfg.get("allow_space_as_indent", False)

        # 字号（用于将"字符"换算为 twips）
        font_cfg = self.config.get("font_check", {})
        font_size_pt = font_cfg.get("size_pt", font_cfg.get("target_size_pt", 14))

        # 换算基准：1 pt = 20 twips；1 字符宽度 ≈ 字号 pt（全角字符近似）
        twips_per_pt = 20
        twips_per_char = font_size_pt * twips_per_pt
        target_indent_twips = target_indent_chars * twips_per_char
        indent_tol_twips = max(100, indent_tol_chars * twips_per_char)

        target_line_twips = target_line_spacing_pt * twips_per_pt
        line_spacing_tol_twips = line_spacing_tol_pt * twips_per_pt

        space_tol_twips = space_tol_pt * twips_per_pt
        target_space_before_twips = space_before_pt * twips_per_pt
        target_space_after_twips = space_after_pt * twips_per_pt

        for i, para in enumerate(doc.paragraphs):
            text_stripped = para.text.strip()
            is_empty = not text_stripped

            # --- 3.0 空段落检查（不得使用回车创建空白行）---
            if is_empty and cfg.get("no_empty_paragraphs", False):
                # 豁免文档最后一个段落（Word 文档结构末尾标记）
                if i < len(doc.paragraphs) - 1:
                    self.errors.append(
                        f"段落 {i + 1}: 发现空段落，不得使用回车创建空白行"
                    )
                continue

            if is_empty:
                continue

            pPr = para._p.pPr

            # --- 3.1 对齐方式 ---
            if pPr is not None and pPr.jc is not None:
                actual_align = self._xml_attr(pPr.jc, "val")
                # Word XML 中 justify 对应 "both"
                expected_align = "both" if target_align == "justify" else target_align
                if actual_align is not None and actual_align != expected_align:
                    self.errors.append(
                        f"段落 {i + 1}: 对齐方式错误 (当前: {actual_align}, 应为: {target_align})"
                    )

            # --- 3.2 行距模式与数值 ---
            if pPr is not None and pPr.spacing is not None:
                spacing = pPr.spacing
                line_rule = self._xml_attr(spacing, "lineRule")
                line_val = self._xml_attr_twips(spacing, "line")

                if line_rule is not None:
                    if line_rule != "exact":
                        self.errors.append(
                            f"段落 {i + 1}: 行距模式错误 (当前: {line_rule}, 应为: exact 固定值)"
                        )
                    elif line_val is not None:
                        if abs(line_val - target_line_twips) > line_spacing_tol_twips:
                            actual_pt = line_val / twips_per_pt
                            self.errors.append(
                                f"段落 {i + 1}: 行距数值错误 (当前: {actual_pt:.1f} pt, 要求: {target_line_spacing_pt} pt 固定值)"
                            )

            # --- 3.3 首行缩进（终极修复版：段落直接属性 + 样式继承链 + 文档默认）---
            has_indent = False
            actual_indent_chars = 0.0
            ind = None  # 提前定义，供 3.3b 使用

            def _read_ind_chars_from_pPr(pPr_el, twips_per_char):
                """从 w:pPr 下的 w:ind 读取首行缩进字符数"""
                if pPr_el is None:
                    return None
                w_ind = pPr_el.find(qn('w:ind'))
                if w_ind is None:
                    return None
                # 1. 优先读 firstLineChars（如 200 = 2字符）
                flc = w_ind.get(qn('w:firstLineChars'))
                if flc is not None:
                    try:
                        return int(flc) / 100
                    except (ValueError, TypeError):
                        pass
                # 2. 再读 firstLine（twips）
                fl = w_ind.get(qn('w:firstLine'))
                if fl is not None:
                    try:
                        return int(fl) / twips_per_char
                    except (ValueError, TypeError):
                        pass
                return None

            # 尝试 1：段落直接属性
            if pPr is not None:
                ind = pPr.ind  # 保留原变量供 3.3b 使用
                val = _read_ind_chars_from_pPr(pPr, twips_per_char)
                if val is not None:
                    actual_indent_chars = val
                    has_indent = True

            # 尝试 2：段落样式链（当前样式 -> 基样式 -> ...）
            if not has_indent:
                try:
                    style = para.style
                    if style and style.element is not None:
                        # 当前样式
                        val = _read_ind_chars_from_pPr(
                            style.element.find(qn('w:pPr')), twips_per_char
                        )
                        if val is not None:
                            actual_indent_chars = val
                            has_indent = True
                        else:
                            # 递归查找基样式（basedOn）
                            doc = para.part.document
                            styles_el = doc.styles.element
                            current_style_el = style.element
                            visited = set()
                            while not has_indent and current_style_el is not None:
                                style_id = current_style_el.get(qn('w:styleId'))
                                if style_id in visited:
                                    break
                                visited.add(style_id)

                                based_on_id = current_style_el.get(qn('w:basedOn'))
                                if not based_on_id:
                                    break

                                # 在样式集中查找基样式
                                xpath = f'.//{qn("w:style")}[@{qn("w:styleId")}="{based_on_id}"]'
                                base_el = styles_el.find(xpath)
                                if base_el is None:
                                    break

                                val = _read_ind_chars_from_pPr(
                                    base_el.find(qn('w:pPr')), twips_per_char
                                )
                                if val is not None:
                                    actual_indent_chars = val
                                    has_indent = True
                                    break

                                current_style_el = base_el
                except Exception:
                    pass

            # 尝试 3：文档默认属性（docDefaults）
            if not has_indent:
                try:
                    doc = para.part.document
                    styles_el = doc.styles.element
                    doc_defaults = styles_el.find(qn('w:docDefaults'))
                    if doc_defaults is not None:
                        pPr_default = doc_defaults.find(qn('w:pPrDefault'))
                        if pPr_default is not None:
                            val = _read_ind_chars_from_pPr(
                                pPr_default.find(qn('w:pPr')), twips_per_char
                            )
                            if val is not None:
                                actual_indent_chars = val
                                has_indent = True
                except Exception:
                    pass

            # 尝试 4：python-docx 原生解析（兜底）
            if not has_indent:
                try:
                    inherited = para.paragraph_format.first_line_indent
                    if inherited is not None:
                        actual_indent_chars = inherited.pt / font_size_pt
                        has_indent = True
                except Exception:
                    pass

            # 判断与报错
            if not has_indent:
                if target_indent_chars > 0:
                    self.errors.append(
                        f"段落 {i + 1}: 缺少首行缩进设置 (要求: {target_indent_chars} 字符)"
                    )
            else:
                if abs(actual_indent_chars - target_indent_chars) > indent_tol_chars:
                    self.errors.append(
                        f"段落 {i + 1}: 首行缩进错误 "
                        f"(当前约 {actual_indent_chars:.1f} 字符, 要求: {target_indent_chars} 字符)"
                    )

            # --- 3.3b 左右缩进检查（不得用缩进增加间距）---
            if not cfg.get("allow_left_indent", True):
                left = self._xml_attr_twips(ind, "left") if ind is not None else None
                if left is not None and left > 50:
                    self.errors.append(
                        f"段落 {i + 1}: 发现左缩进设置，暗标不得使用左缩进增加间距"
                    )
            if not cfg.get("allow_right_indent", True):
                right = self._xml_attr_twips(ind, "right") if ind is not None else None
                if right is not None and right > 50:
                    self.errors.append(
                        f"段落 {i + 1}: 发现右缩进设置，暗标不得使用右缩进增加间距"
                    )

            # --- 3.4 空格/Tab 代替缩进 ---
            if not allow_space_as_indent:
                text = para.text
                if text and (text.startswith(" ") or text.startswith("\t")):
                    self.errors.append(
                        f"段落 {i + 1}: 发现使用空格或 Tab 代替缩进，请使用段落格式设置"
                    )

            # --- 3.5 段前段后间距 ---
            if pPr is not None and pPr.spacing is not None:
                before = self._xml_attr_twips(pPr.spacing, "before")
                after = self._xml_attr_twips(pPr.spacing, "after")
                before_twips = before if before is not None else 0
                after_twips = after if after is not None else 0

                if abs(before_twips - target_space_before_twips) > space_tol_twips:
                    self.errors.append(
                        f"段落 {i + 1}: 段前间距错误 (当前: {before_twips / twips_per_pt:.1f} pt, 要求: {space_before_pt} pt)"
                    )
                if abs(after_twips - target_space_after_twips) > space_tol_twips:
                    self.errors.append(
                        f"段落 {i + 1}: 段后间距错误 (当前: {after_twips / twips_per_pt:.1f} pt, 要求: {space_after_pt} pt)"
                    )

    # ------------------------------------------------------------------
    # 4. 表格与图片
    # ------------------------------------------------------------------
    def _check_objects(self, doc):
        cfg = self.config.get("object_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查表格和图片格式...")

        # --- 表格 ---
        for i, table in enumerate(doc.tables):
            # 表格整体居中
            if cfg.get("table_alignment") == "center":
                if table.alignment is not None and table.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.errors.append(f"表格 {i + 1}: 表格整体未居中对齐")

            # 遍历表格内的段落
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        p_fmt = para.paragraph_format
                        p_pr = para._element.pPr

                        # 表内文字左对齐
                        if cfg.get("table_text_alignment") == "left":
                            if p_pr is not None and p_pr.jc is not None:
                                actual = self._xml_attr(p_pr.jc, "val")
                                if actual is not None and actual != "left":
                                    self.errors.append(f"表格 {i + 1} 单元格: 表内文字未左对齐 (当前: {actual})")

                        # 表内无首行缩进
                        if cfg.get("table_indent_none"):
                            if p_pr is not None and p_pr.ind is not None:
                                first_line = self._xml_attr_twips(p_pr.ind, "firstLine")
                                if first_line and first_line > 50:
                                    self.errors.append(f"表格 {i + 1} 单元格: 表内文字不应有首行缩进")

                        # 表内行距不能为固定值
                        if cfg.get("table_spacing_none"):
                            if p_pr is not None and p_pr.spacing is not None:
                                line_rule = self._xml_attr(p_pr.spacing, "lineRule")
                                if line_rule == "exact":
                                    self.errors.append(f"表格 {i + 1} 单元格: 表内文字行距不能为固定值")

        # --- 图片 ---
        for i, para in enumerate(doc.paragraphs):
            if "w:drawing" in para._p.xml or "w:pict" in para._p.xml:
                if cfg.get("image_alignment") == "center":
                    pPr = para._p.pPr
                    if pPr is not None and pPr.jc is not None:
                        actual = self._xml_attr(pPr.jc, "val")
                        if actual is not None and actual != "center":
                            self.errors.append(f"段落 {i + 1}: 图片未居中对齐 (当前: {actual})")

    # ------------------------------------------------------------------
    # 5. 文档结构（页眉页脚页码）
    # ------------------------------------------------------------------
    def _check_structure(self, doc):
        cfg = self.config.get("structure_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查文档结构...")
        for i, section in enumerate(doc.sections):
            # 检查页眉
            if not cfg.get("allow_header", False):
                header_text = "".join([p.text.strip() for p in section.header.paragraphs])
                if header_text:
                    self.errors.append(f"第 {i + 1} 节：发现违规页眉内容 -> '{header_text[:20]}'")

            # 检查页脚（普通文本内容）
            if not cfg.get("allow_footer", False):
                footer_text = "".join([p.text.strip() for p in section.footer.paragraphs])
                if footer_text:
                    self.errors.append(f"第 {i + 1} 节：发现违规页脚内容 -> '{footer_text[:20]}'")

            # 检查页码（域代码，需要检查XML）
            if not cfg.get("allow_page_number", False):
                footer = section.footer
                # 检查页脚XML中是否有页码域代码 (w:fldChar 或 w:fldSimple)
                footer_xml = footer._element.xml if footer._element is not None else ""
                has_page_number = "w:fldChar" in footer_xml or "w:fldSimple" in footer_xml or "PAGE" in footer_xml

                # 同时检查页眉中是否有页码
                header = section.header
                header_xml = header._element.xml if header._element is not None else ""
                has_page_number_in_header = "w:fldChar" in header_xml or "w:fldSimple" in header_xml or "PAGE" in header_xml

                if has_page_number or has_page_number_in_header:
                    location = "页眉" if has_page_number_in_header else "页脚"
                    self.errors.append(f"第 {i + 1} 节：发现违规页码（{location}中检测到页码域代码），不得设置页码")

    # ------------------------------------------------------------------
    # 6. 标点符号检查
    # ------------------------------------------------------------------
    def _check_punctuation(self, doc):
        cfg = self.config.get("punctuation_check", {})
        if not cfg.get("enabled"):
            return

        log("正在检查标点符号...")
        # 常见英文标点，在中文字符之间出现时视为疑似错误
        pattern = re.compile(r"(?<=[\u4e00-\u9fa5])[.,;:!?](?=[\u4e00-\u9fa5])")

        for i, para in enumerate(doc.paragraphs):
            text = para.text
            if not text.strip():
                continue
            matches = pattern.findall(text)
            if matches:
                unique = sorted(set(matches))
                display = ", ".join(unique[:5])
                self.errors.append(
                    f"段落 {i + 1}: 发现英文标点混用 -> [{display}]，请使用中文标点"
                )

    def report_results(self):
        """
        打印最终检查结果
        """
        print("\n" + "=" * 30)
        if not self.errors:
            print("✅ 检查通过！文档符合暗标排版要求。")
        else:
            print(f"❌ 检查完成，发现 {len(self.errors)} 处问题：")
            for err in self.errors:
                print(f" - {err}")
        print("=" * 30)