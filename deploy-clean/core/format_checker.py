"""
format_checker.py - Word 暗标格式检查器（规则驱动版）
支持：页面、结构、边距、字体、段落、表格（含垂直对齐）、标题、标点等全面检查。
"""

import json
import re
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


def log(msg):
    print(f"[检查中] {msg}")


class FormatChecker:
    def __init__(self, config_path="config/rules.json"):
        with open(config_path, "r", encoding="utf-8") as f:
            self.config = json.load(f)
        self.errors = []
        self.rule_summary = []
        self.punctuation_errors = []

        # 表格专用配置（从 table_check 读取）
        table_cfg = self.config.get("table_check", {})
        self.table_font_enabled = table_cfg.get("enabled", False)
        self.table_font_family = table_cfg.get("table_font_family")  # 默认 None 表示使用全局
        self.table_font_size_pt = table_cfg.get("table_font_size_pt")
        self.table_font_color_hex = table_cfg.get("table_font_color_hex")
        self.table_allow_bold = table_cfg.get("table_allow_bold")
        self.table_allow_italic = table_cfg.get("table_allow_italic")
        self.table_allow_underline = table_cfg.get("table_allow_underline")
        self.table_allow_color_change = table_cfg.get("table_allow_color_change")
        self.table_vertical_alignment = table_cfg.get("table_vertical_alignment")  # 新增

        # 标题检查配置
        self.heading_cfg = self.config.get("heading_check", {})

    def _generate_rule_summary(self):
        """生成规则清单"""
        summary = []
        summary.append("========== 暗标检测规则清单 ==========")

        info = self.config.get("document_info", {})
        summary.append(f"规则名称: {info.get('name', '未命名')}")
        summary.append(f"规则描述: {info.get('description', '无')}")
        summary.append(f"生成日期: {info.get('generated_date', '未知')}")

        # 页面检查
        page = self.config.get("page_check", {})
        if page.get("enabled") and page.get("paper_size"):
            summary.append("✅ 页面检查: 已启用 (纸张A4, 页数限制, 禁止空白页)")
        else:
            summary.append("❌ 页面检查: 未启用或未完整配置")

        # 结构检查
        struct = self.config.get("structure_check", {})
        if struct.get("enabled"):
            items = []
            if struct.get("allow_toc") is not None:
                items.append("目录" if struct["allow_toc"] else "禁止目录")
            if struct.get("allow_header") is not None:
                items.append("页眉" if struct["allow_header"] else "禁止页眉")
            if struct.get("allow_footer") is not None:
                items.append("页脚" if struct["allow_footer"] else "禁止页脚")
            if struct.get("allow_page_number") is not None:
                items.append("页码" if struct["allow_page_number"] else "禁止页码")
            if struct.get("allow_cover") is not None:
                items.append("封面" if struct["allow_cover"] else "禁止封面")
            summary.append("✅ 结构检查: 已启用 (" + ", ".join(items) + ")")
        else:
            summary.append("❌ 结构检查: 未启用")

        # 页边距
        margin = self.config.get("margin_check", {})
        if margin.get("enabled") and all(margin.get(k) is not None for k in ["top_cm", "bottom_cm", "left_cm", "right_cm"]):
            summary.append(f"✅ 页边距检查: 已启用 (上{margin['top_cm']}cm, 下{margin['bottom_cm']}cm, 左{margin['left_cm']}cm, 右{margin['right_cm']}cm)")
        else:
            summary.append("❌ 页边距检查: 未启用或数值不完整")

        # 字体
        font = self.config.get("font_check", {})
        if font.get("enabled"):
            sum_str = f"✅ 字体检查: 已启用 (中文字体={font.get('chinese_font','未指定')}, 字号={font.get('size_pt')}pt, 黑色, 禁止加粗/倾斜/下划线/变色)"
            summary.append(sum_str)
        else:
            summary.append("❌ 字体检查: 未启用")

        # 段落
        para = self.config.get("paragraph_check", {})
        if para.get("enabled"):
            items = []
            if para.get("alignment"):
                items.append(f"对齐={para['alignment']}")
            if para.get("line_spacing_rule") == "exact" and para.get("line_spacing_pt"):
                items.append(f"行距={para['line_spacing_pt']}pt固定值")
            if para.get("first_line_indent_chars"):
                items.append(f"首行缩进={para['first_line_indent_chars']}字符")
            if para.get("space_before_pt") is not None:
                items.append(f"段前={para['space_before_pt']}pt")
            if para.get("space_after_pt") is not None:
                items.append(f"段后={para['space_after_pt']}pt")
            if para.get("allow_space_as_indent") is False:
                items.append("禁用空格/Tab缩进")
            if para.get("allow_left_indent") is False:
                items.append("禁用左缩进")
            if para.get("allow_right_indent") is False:
                items.append("禁用右缩进")
            if para.get("no_empty_paragraphs"):
                items.append("禁止空段落")
            if items:
                summary.append("✅ 段落检查: 已启用 (" + ", ".join(items) + ")")
            else:
                summary.append("✅ 段落检查: 已启用 (未配置具体参数)")
        else:
            summary.append("❌ 段落检查: 未启用")

        # 表格
        table = self.config.get("table_check", {})
        if table.get("enabled"):
            items = []
            if table.get("table_alignment"):
                items.append(f"表格对齐={table['table_alignment']}")
            if table.get("table_vertical_alignment"):
                items.append(f"表内垂直对齐={table['table_vertical_alignment']}")
            if table.get("table_text_alignment"):
                items.append(f"表内文字对齐={table['table_text_alignment']}")
            if table.get("table_indent_none"):
                items.append("表内无缩进")
            if table.get("table_spacing_none"):
                items.append("表内非固定行距")
            if table.get("table_font_family") or table.get("table_font_size_pt"):
                items.append(f"表内字体={table.get('table_font_family','未指定')} {table.get('table_font_size_pt','?')}pt")
            summary.append("✅ 表格检查: 已启用 (" + ", ".join(items) + ")")
        else:
            summary.append("❌ 表格检查: 未启用")

        # 标题检查
        heading = self.config.get("heading_check", {})
        if heading.get("enabled"):
            levels = heading.get("level_rules", {})
            level_desc = ", ".join([f"{k}级" for k in levels.keys() if levels[k].get("enabled")])
            summary.append(f"✅ 标题检查: 已启用 (级别: {level_desc})")
        else:
            summary.append("❌ 标题检查: 未启用")

        # 标点
        punct = self.config.get("punctuation_check", {})
        if punct.get("enabled"):
            items = []
            if punct.get("require_chinese"):
                items.append("中文标点")
            if punct.get("no_space_between_chars"):
                items.append("禁止字符间空格")
            summary.append("✅ 标点检查: 已启用 (" + ", ".join(items) + ")")
        else:
            summary.append("❌ 标点检查: 未启用")

        summary.append("=====================================")
        self.rule_summary = summary
        for line in summary:
            log(line)

    def check_document(self, file_path):
        self.errors = []
        self.punctuation_errors = []
        self._generate_rule_summary()
        log(f"正在加载文档: {file_path}")
        try:
            doc = Document(file_path)
        except Exception as e:
            self.errors.append(f"无法打开文档: {e}")
            return self.errors

        # 按顺序检查
        self._check_page_setup(doc)
        self._check_structure(doc)
        self._check_margins(doc)
        self._check_fonts_and_styles(doc)
        self._check_paragraph_format(doc)
        self._check_tables(doc)
        self._check_headings(doc)
        self._check_punctuation(doc)

        return self.errors

    # ------------------- 辅助函数 -------------------
    @staticmethod
    def _xml_attr(element, attr_local_name):
        if element is None:
            return None
        return element.get(qn(f"w:{attr_local_name}"))

    @staticmethod
    def _xml_attr_twips(element, attr_local_name):
        val = FormatChecker._xml_attr(element, attr_local_name)
        if val is None:
            return None
        try:
            return int(val)
        except (ValueError, TypeError):
            return None

    def _get_cell_vertical_alignment(self, cell):
        """获取单元格垂直对齐方式，返回 'top'/'center'/'bottom'，若无则返回 None"""
        tc = cell._element
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is not None:
            vAlign = tcPr.find(qn('w:vAlign'))
            if vAlign is not None:
                return vAlign.get(qn('w:val'))
        return None

    def _check_run_font(self, run, para_idx, context="段落",
                        target_font=None, target_size_pt=None,
                        target_color=None, allow_bold=None,
                        allow_italic=None, allow_underline=None,
                        allow_color_change=None):
        """
        检查 run 的字体、字号、颜色、加粗、倾斜、下划线。
        若参数为 None，则跳过对应检查（表示不要求）。
        """
        cfg = self.config.get("font_check", {})
        if not cfg.get("enabled"):
            return

        text = run.text
        if not text.strip():
            return

        # 确定检查目标值：优先使用传入参数，否则从配置读取（可能为 None）
        if target_font is None:
            target_font = cfg.get("chinese_font")
        if target_size_pt is None:
            target_size_pt = cfg.get("size_pt")
        if target_color is None:
            target_color = cfg.get("color_hex", "000000")
        if allow_bold is None:
            allow_bold = cfg.get("allow_bold", False)
        if allow_italic is None:
            allow_italic = cfg.get("allow_italic", False)
        if allow_underline is None:
            allow_underline = cfg.get("allow_underline", False)
        if allow_color_change is None:
            allow_color_change = cfg.get("allow_color_change", False)

        # 1. 字体检查（仅当 target_font 非 None）
        if target_font is not None:
            rpr = run._element.rPr
            if rpr is not None:
                rfonts = rpr.rFonts
                if rfonts is not None:
                    east_asia = rfonts.get(qn("w:eastAsia"))
                    if east_asia and target_font not in east_asia:
                        self.errors.append(
                            f"{context} {para_idx}: 字体错误 '{text[:10]}...' (当前: {east_asia}, 应为: {target_font})"
                        )
                    elif not east_asia:
                        if run.font.name and target_font not in run.font.name:
                            self.errors.append(
                                f"{context} {para_idx}: 字体错误 '{text[:10]}...' (当前: {run.font.name}, 应为: {target_font})"
                            )
                else:
                    if run.font.name and target_font not in run.font.name:
                        self.errors.append(
                            f"{context} {para_idx}: 字体错误 '{text[:10]}...' (当前: {run.font.name}, 应为: {target_font})"
                        )
            else:
                if run.font.name and target_font not in run.font.name:
                    self.errors.append(
                        f"{context} {para_idx}: 字体错误 '{text[:10]}...' (当前: {run.font.name}, 应为: {target_font})"
                    )

        # 2. 字号检查（仅当 target_size_pt 非 None）
        if target_size_pt is not None and run.font.size is not None:
            actual_pt = run.font.size.pt
            if abs(actual_pt - target_size_pt) > 0.5:
                self.errors.append(
                    f"{context} {para_idx}: 字号错误 (当前: {actual_pt} pt, 应为: {target_size_pt} pt)"
                )

        # 3. 颜色检查（仅当 target_color 非 None 且 allow_color_change 为 False 时）
        if target_color is not None and allow_color_change is False:
            if run.font.color and run.font.color.rgb:
                color_hex = str(run.font.color.rgb)
                if color_hex.upper() != target_color.upper():
                    self.errors.append(
                        f"{context} {para_idx}: 字体颜色错误 (当前: {color_hex}, 必须为黑色 {target_color})"
                    )

        # 4. 加粗/倾斜/下划线（仅当对应允许标志为 False 时检查）
        if allow_bold is False and run.bold:
            self.errors.append(
                f"{context} {para_idx}: 发现违规加粗内容 -> '{text.strip()[:10]}'"
            )
        if allow_italic is False and run.italic:
            self.errors.append(
                f"{context} {para_idx}: 发现违规倾斜内容 -> '{text.strip()[:10]}'"
            )
        if allow_underline is False and run.underline:
            self.errors.append(
                f"{context} {para_idx}: 发现违规下划线内容 -> '{text.strip()[:10]}'"
            )

    # ------------------- 各项检查 -------------------
    def _check_page_setup(self, doc):
        cfg = self.config.get("page_check", {})
        if not cfg.get("enabled"):
            return
        log("检查纸张大小、页数限制及空白页...")
        max_pages = cfg.get("max_pages")
        if max_pages is not None:
            para_count = len([p for p in doc.paragraphs if p.text.strip()])
            estimated_pages = para_count // 35 + 1
            if estimated_pages > max_pages:
                self.errors.append(
                    f"文档页数可能超标 (估算约 {estimated_pages} 页，限制 {max_pages} 页)"
                )
        target_paper = cfg.get("paper_size")
        if target_paper == "A4":
            for i, section in enumerate(doc.sections):
                w_cm = section.page_width.cm
                h_cm = section.page_height.cm
                is_a4 = (abs(w_cm - 21.0) < 0.5 and abs(h_cm - 29.7) < 0.5) or \
                        (abs(w_cm - 29.7) < 0.5 and abs(h_cm - 21.0) < 0.5)
                if not is_a4:
                    self.errors.append(
                        f"第 {i+1} 节：纸张大小错误 (当前: {w_cm:.1f}×{h_cm:.1f} cm)，必须为 A4"
                    )
        # 空白页检查
        if cfg.get("no_blank_pages", False):
            for i, section in enumerate(doc.sections):
                # 简单判断：检查该节是否有任何内容（段落或表格）
                has_content = False
                # 检查页眉页脚内容（通常不应有）
                for para in section.header.paragraphs + section.footer.paragraphs:
                    if para.text.strip():
                        has_content = True
                        break
                if has_content:
                    continue
                # 检查正文段落
                for para in doc.paragraphs:
                    if para.text.strip():
                        has_content = True
                        break
                if not has_content:
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    has_content = True
                                    break
                            if has_content:
                                break
                        if has_content:
                            break
                if not has_content:
                    self.errors.append(f"第 {i+1} 节：发现空白页，暗标不允许出现空白页。")

    def _check_structure(self, doc):
        cfg = self.config.get("structure_check", {})
        if not cfg.get("enabled"):
            return
        log("检查文档结构...")
        for i, section in enumerate(doc.sections):
            if cfg.get("allow_header") is False:
                header_text = "".join(p.text.strip() for p in section.header.paragraphs)
                if header_text:
                    self.errors.append(f"第 {i+1} 节：发现违规页眉内容 -> '{header_text[:20]}'")
            if cfg.get("allow_footer") is False:
                footer_text = "".join(p.text.strip() for p in section.footer.paragraphs)
                if footer_text:
                    self.errors.append(f"第 {i+1} 节：发现违规页脚内容 -> '{footer_text[:20]}'")
            if cfg.get("allow_page_number") is False:
                footer_xml = section.footer._element.xml if section.footer._element is not None else ""
                header_xml = section.header._element.xml if section.header._element is not None else ""
                has_page = ("w:fldChar" in footer_xml or "w:fldSimple" in footer_xml or "PAGE" in footer_xml or
                            "w:fldChar" in header_xml or "w:fldSimple" in header_xml or "PAGE" in header_xml)
                if has_page:
                    location = "页眉" if ("w:fldChar" in header_xml or "PAGE" in header_xml) else "页脚"
                    self.errors.append(f"第 {i+1} 节：发现违规页码（{location}中检测到页码域代码）")

    def _check_margins(self, doc):
        cfg = self.config.get("margin_check", {})
        if not cfg.get("enabled"):
            return
        top = cfg.get("top_cm")
        bottom = cfg.get("bottom_cm")
        left = cfg.get("left_cm")
        right = cfg.get("right_cm")
        if any(v is None for v in [top, bottom, left, right]):
            log("页边距配置不完整，跳过检查")
            return
        tolerance = cfg.get("tolerance_cm", 0.1)
        log("检查页边距...")
        for i, section in enumerate(doc.sections):
            actuals = {
                "top": section.top_margin.cm,
                "bottom": section.bottom_margin.cm,
                "left": section.left_margin.cm,
                "right": section.right_margin.cm,
            }
            for name, target in {"top": top, "bottom": bottom, "left": left, "right": right}.items():
                actual = actuals[name]
                if abs(actual - target) > tolerance:
                    self.errors.append(
                        f"第 {i+1} 节：{name}边距错误 (当前: {actual:.2f} cm, 要求: {target} cm)"
                    )

    def _check_fonts_and_styles(self, doc):
        cfg = self.config.get("font_check", {})
        if not cfg.get("enabled"):
            return
        log("检查字体、字号及样式...")
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                continue
            for run in para.runs:
                self._check_run_font(run, i+1, "段落")

    def _check_paragraph_format(self, doc):
        cfg = self.config.get("paragraph_check", {})
        if not cfg.get("enabled"):
            return
        log("检查段落格式...")

        target_align = cfg.get("alignment")
        target_line_spacing_pt = cfg.get("line_spacing_pt")
        line_spacing_tol_pt = cfg.get("line_spacing_tolerance_pt", 0.5)
        target_indent_chars = cfg.get("first_line_indent_chars")
        indent_tol_chars = cfg.get("first_line_indent_tolerance_chars", 0.3)
        space_before_pt = cfg.get("space_before_pt")
        space_after_pt = cfg.get("space_after_pt")
        space_tol_pt = cfg.get("space_tolerance_pt", 0.5)
        allow_space_as_indent = cfg.get("allow_space_as_indent")
        no_empty_paragraphs = cfg.get("no_empty_paragraphs", False)
        allow_left_indent = cfg.get("allow_left_indent")
        allow_right_indent = cfg.get("allow_right_indent")

        font_cfg = self.config.get("font_check", {})
        font_size_pt = font_cfg.get("size_pt", 14)
        twips_per_pt = 20

        target_line_twips = target_line_spacing_pt * twips_per_pt if target_line_spacing_pt is not None else None
        line_spacing_tol_twips = line_spacing_tol_pt * twips_per_pt
        space_tol_twips = space_tol_pt * twips_per_pt
        target_space_before_twips = space_before_pt * twips_per_pt if space_before_pt is not None else None
        target_space_after_twips = space_after_pt * twips_per_pt if space_after_pt is not None else None

        for i, para in enumerate(doc.paragraphs):
            text_stripped = para.text.strip()
            is_empty = not text_stripped

            if no_empty_paragraphs and is_empty:
                if i < len(doc.paragraphs) - 1:
                    self.errors.append(f"段落 {i+1}: 发现空段落，不得使用回车创建空白行")
                continue
            if is_empty:
                continue

            pPr = para._p.pPr

            # ---------- 对齐方式检查（优化版） ----------
            if target_align is not None:
                actual_align = None
                # 优先从 XML 读取
                if pPr is not None and pPr.jc is not None:
                    actual_align = self._xml_attr(pPr.jc, "val")
                    if actual_align == "both":
                        actual_align = "justify"
                # 若 XML 无值，尝试从 paragraph_format 获取
                if actual_align is None:
                    try:
                        para_align = para.paragraph_format.alignment
                        if para_align is not None:
                            if para_align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                actual_align = "justify"
                            elif para_align == WD_ALIGN_PARAGRAPH.LEFT:
                                actual_align = "left"
                            elif para_align == WD_ALIGN_PARAGRAPH.CENTER:
                                actual_align = "center"
                            elif para_align == WD_ALIGN_PARAGRAPH.RIGHT:
                                actual_align = "right"
                    except Exception:
                        pass
                # 如果仍无法获取，且要求左对齐，则视为符合（Word 默认左对齐）
                if actual_align is not None and actual_align != target_align:
                    self.errors.append(f"段落 {i + 1}: 对齐方式错误 (当前: {actual_align}, 应为: {target_align})")

            # 行距
            if target_line_twips is not None and pPr is not None and pPr.spacing is not None:
                spacing = pPr.spacing
                line_rule = self._xml_attr(spacing, "lineRule")
                line_val = self._xml_attr_twips(spacing, "line")
                if line_rule is not None:
                    if line_rule != "exact":
                        self.errors.append(f"段落 {i+1}: 行距模式错误 (当前: {line_rule}, 应为: exact 固定值)")
                    elif line_val is not None:
                        if abs(line_val - target_line_twips) > line_spacing_tol_twips:
                            actual_pt = line_val / twips_per_pt
                            self.errors.append(f"段落 {i+1}: 行距数值错误 (当前: {actual_pt:.1f} pt, 要求: {target_line_spacing_pt} pt 固定值)")

            # 首行缩进
            if target_indent_chars is not None:
                actual_indent_chars = None
                try:
                    first_line_indent = para.paragraph_format.first_line_indent
                    if first_line_indent is not None:
                        indent_pt = first_line_indent.pt
                        if indent_pt is not None and indent_pt > 0:
                            actual_indent_chars = indent_pt / font_size_pt
                except Exception:
                    pass
                if actual_indent_chars is None and pPr is not None:
                    try:
                        ind = pPr.ind
                        if ind is not None:
                            first_line_chars = ind.get(qn('w:firstLineChars'))
                            if first_line_chars is not None:
                                actual_indent_chars = int(first_line_chars) / 100.0
                            else:
                                first_line = ind.get(qn('w:firstLine'))
                                if first_line is not None:
                                    indent_twips = int(first_line)
                                    indent_pt = indent_twips / 20.0
                                    actual_indent_chars = indent_pt / font_size_pt
                    except Exception:
                        pass
                if actual_indent_chars is not None and abs(actual_indent_chars - target_indent_chars) > indent_tol_chars:
                    self.errors.append(
                        f"段落 {i + 1}: 首行缩进错误 (当前约 {actual_indent_chars:.1f} 字符, 要求: {target_indent_chars} 字符)")

            # 左右缩进禁用
            if pPr is not None and pPr.ind is not None:
                ind = pPr.ind
                if allow_left_indent is False:
                    left_twips = self._xml_attr_twips(ind, "left")
                    if left_twips is not None and left_twips > 50:
                        self.errors.append(f"段落 {i+1}: 发现左缩进设置，暗标不得使用左缩进")
                if allow_right_indent is False:
                    right_twips = self._xml_attr_twips(ind, "right")
                    if right_twips is not None and right_twips > 50:
                        self.errors.append(f"段落 {i+1}: 发现右缩进设置，暗标不得使用右缩进")

            # 空格/Tab缩进
            if allow_space_as_indent is False:
                if text_stripped and (text_stripped.startswith(" ") or text_stripped.startswith("\t")):
                    self.errors.append(f"段落 {i+1}: 发现使用空格或 Tab 代替缩进，请使用段落格式设置")

            # 段前段后间距
            if pPr is not None and pPr.spacing is not None:
                spacing = pPr.spacing
                before = self._xml_attr_twips(spacing, "before")
                after = self._xml_attr_twips(spacing, "after")
                before_twips = before if before is not None else 0
                after_twips = after if after is not None else 0
                if target_space_before_twips is not None and abs(before_twips - target_space_before_twips) > space_tol_twips:
                    self.errors.append(f"段落 {i+1}: 段前间距错误 (当前: {before_twips / twips_per_pt:.1f} pt, 要求: {space_before_pt} pt)")
                if target_space_after_twips is not None and abs(after_twips - target_space_after_twips) > space_tol_twips:
                    self.errors.append(f"段落 {i+1}: 段后间距错误 (当前: {after_twips / twips_per_pt:.1f} pt, 要求: {space_after_pt} pt)")

    def _check_tables(self, doc):
        """检查表格（不检查图片）"""
        cfg = self.config.get("table_check", {})
        if not cfg.get("enabled"):
            return
        log("检查表格...")

        for i, table in enumerate(doc.tables):
            if cfg.get("table_alignment") == "center":
                if table.alignment is not None and table.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                    self.errors.append(f"表格 {i+1}: 表格整体未居中对齐")

            for row in table.rows:
                for cell in row.cells:
                    # ---------- 垂直对齐检查 ----------
                    if self.table_vertical_alignment is not None:
                        actual_valign = self._get_cell_vertical_alignment(cell)
                        if actual_valign is None:
                            self.errors.append(f"表格 {i+1} 单元格: 未设置垂直对齐方式 (要求 {self.table_vertical_alignment})")
                        elif actual_valign != self.table_vertical_alignment:
                            self.errors.append(f"表格 {i+1} 单元格: 垂直对齐错误 (当前: {actual_valign}, 要求: {self.table_vertical_alignment})")

                    for para in cell.paragraphs:
                        if not para.text.strip():
                            continue
                        p_pr = para._element.pPr

                        if cfg.get("table_text_alignment") == "left":
                            if p_pr is not None and p_pr.jc is not None:
                                actual = self._xml_attr(p_pr.jc, "val")
                                if actual is not None and actual != "left":
                                    self.errors.append(f"表格 {i+1} 单元格: 表内文字未左对齐")

                        if cfg.get("table_indent_none"):
                            if p_pr is not None and p_pr.ind is not None:
                                first_line = self._xml_attr_twips(p_pr.ind, "firstLine")
                                if first_line and first_line > 50:
                                    self.errors.append(f"表格 {i+1} 单元格: 表内文字不应有首行缩进")

                        if cfg.get("table_spacing_none"):
                            if p_pr is not None and p_pr.spacing is not None:
                                line_rule = self._xml_attr(p_pr.spacing, "lineRule")
                                if line_rule == "exact":
                                    self.errors.append(f"表格 {i+1} 单元格: 表内文字行距不能为固定值")

                        # 表格字体检查（独立于正文）
                        if self.table_font_enabled:
                            for run in para.runs:
                                self._check_run_font(
                                    run, i+1, context="表格",
                                    target_font=self.table_font_family,
                                    target_size_pt=self.table_font_size_pt,
                                    target_color=self.table_font_color_hex,
                                    allow_bold=self.table_allow_bold,
                                    allow_italic=self.table_allow_italic,
                                    allow_underline=self.table_allow_underline,
                                    allow_color_change=self.table_allow_color_change
                                )

    def _check_headings(self, doc):
        """识别标题并检查格式（字体、字号、对齐、编号后缀等）"""
        if not self.config.get("heading_check", {}).get("enabled", False):
            return
        log("检查标题格式...")

        ident = self.config.get("heading_check", {}).get("identification", {})
        patterns = ident.get("patterns", {})
        level_rules = self.config.get("heading_check", {}).get("level_rules", {})

        # 编译正则表达式
        compiled = {}
        for level, pat in patterns.items():
            try:
                compiled[level] = re.compile(pat)
            except:
                pass

        for idx, para in enumerate(doc.paragraphs, start=1):
            text = para.text.strip()
            if not text:
                continue

            matched_level = None
            for level, pattern in compiled.items():
                if pattern.match(text):
                    matched_level = level
                    break

            if matched_level is None:
                continue

            rule = level_rules.get(matched_level, {})
            if not rule.get("enabled", False):
                continue

            # 检查字体格式
            for run in para.runs:
                self._check_run_font(
                    run,
                    para_idx=idx,
                    context=f"标题 {matched_level}",
                    target_font=rule.get("chinese_font", "宋体"),
                    target_size_pt=rule.get("size_pt"),
                    target_color=rule.get("color_hex", "000000"),
                    allow_bold=rule.get("allow_bold", False),
                    allow_italic=rule.get("allow_italic", False),
                    allow_underline=rule.get("allow_underline", False),
                    allow_color_change=rule.get("allow_color_change", False)
                )

            # 检查对齐方式
            target_align = rule.get("alignment")
            if target_align is not None:
                actual_align = None
                pPr = para._p.pPr
                if pPr is not None and pPr.jc is not None:
                    actual_align = self._xml_attr(pPr.jc, "val")
                    if actual_align == "both":
                        actual_align = "justify"
                if actual_align is None:
                    try:
                        para_align = para.paragraph_format.alignment
                        if para_align is not None:
                            if para_align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                                actual_align = "justify"
                            elif para_align == WD_ALIGN_PARAGRAPH.LEFT:
                                actual_align = "left"
                            elif para_align == WD_ALIGN_PARAGRAPH.CENTER:
                                actual_align = "center"
                            elif para_align == WD_ALIGN_PARAGRAPH.RIGHT:
                                actual_align = "right"
                    except Exception:
                        pass
                if actual_align is not None and actual_align != target_align:
                    self.errors.append(
                        f"标题 {matched_level} (段落 {idx}): 对齐方式错误 (当前: {actual_align}, 应为: {target_align})")

            # ---------- 编号后缀检查（增加防御性判断） ----------
            num_format = rule.get("number_format", {})
            # 如果 number_format 为 None 或空字典，则跳过后缀检查
            if num_format:
                required_suffix = num_format.get("required_suffix")
                allowed_suffix = num_format.get("allowed_suffix", [])
                forbid_suffix = num_format.get("forbid_suffix", [])

                if required_suffix is not None:
                    m = re.match(r'^([\d.]+)(.)', text)
                    if m:
                        suffix = m.group(2)
                        if suffix != required_suffix:
                            self.errors.append(
                                f"标题 {matched_level} (段落 {idx}): 编号后缀应为 '{required_suffix}'，当前为 '{suffix}'")
                    else:
                        self.errors.append(
                            f"标题 {matched_level} (段落 {idx}): 无法识别编号后缀，要求为 '{required_suffix}'")

                if forbid_suffix:
                    m = re.match(r'^([\d.]+)(.)', text)
                    if m:
                        suffix = m.group(2)
                        if suffix in forbid_suffix:
                            self.errors.append(
                                f"标题 {matched_level} (段落 {idx}): 编号后缀 '{suffix}' 被禁止，允许的后缀为 {allowed_suffix}")

                if allowed_suffix:
                    m = re.match(r'^([\d.]+)(.)', text)
                    if m:
                        suffix = m.group(2)
                        if suffix not in allowed_suffix:
                            self.errors.append(
                                f"标题 {matched_level} (段落 {idx}): 编号后缀 '{suffix}' 不在允许列表 {allowed_suffix} 中")
            # 如果 num_format 为 None 或空，则什么都不做

    def _check_punctuation(self, doc):
        """检查标点符号（中文标点强制，并检查中文字符间空格）"""
        cfg = self.config.get("punctuation_check", {})
        if not cfg.get("enabled"):
            return
        log("检查标点符号及字符空格...")

        # 1. 检查英文标点（中文之间）
        if cfg.get("require_chinese", True):
            pattern = re.compile(r"(?<=[\u4e00-\u9fa5])([.,;:!?])(?=[\u4e00-\u9fa5])")
            for i, para in enumerate(doc.paragraphs, start=1):
                text = para.text
                if not text.strip():
                    continue
                for match in pattern.finditer(text):
                    punct = match.group(1)
                    pos = match.start()
                    ctx_start = max(0, pos - 10)
                    ctx_end = min(len(text), pos + 11)
                    context = text[ctx_start:ctx_end].replace('\n', ' ').replace('\r', ' ')
                    self.punctuation_errors.append((i-1, pos, punct, context))
                    self.errors.append(
                        f"段落 {i} (第{pos+1}字符): 发现英文标点 '{punct}'，附近：…{context}…，请使用中文标点"
                    )

        # 2. 检查中文字符间空格
        if cfg.get("no_space_between_chars", True):
            pattern = re.compile(r'([\u4e00-\u9fa5])[ \u00a0]([\u4e00-\u9fa5])')
            for i, para in enumerate(doc.paragraphs, start=1):
                text = para.text
                if not text.strip():
                    continue
                matches = list(pattern.finditer(text))
                if matches:
                    positions = []
                    for m in matches:
                        start = max(0, m.start() - 2)
                        end = min(len(text), m.end() + 2)
                        positions.append(text[start:end])
                    unique = list(dict.fromkeys(positions))[:3]
                    snippets_str = "; ".join([f"'{s}'" for s in unique])
                    self.errors.append(
                        f"段落 {i}: 发现中文字符间有空格 -> {snippets_str}，请删除"
                    )

    def report_results(self):
        print("\n" + "=" * 30)
        if not self.errors:
            print("✅ 检查通过！文档符合暗标排版要求。")
        else:
            print(f"❌ 检查完成，发现 {len(self.errors)} 处问题：")
            for err in self.errors:
                print(f" - {err}")
        print("=" * 30)