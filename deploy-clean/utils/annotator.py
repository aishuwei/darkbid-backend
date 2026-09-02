"""
utils/annotator.py - 批注版文档生成器（增强版：支持精确位置批注）
职责：将检查结果以 Word 批注形式插入原文对应位置

依赖：python-docx >= 1.2.0
"""
import os
import re
import shutil
from docx import Document


def log(msg):
    print(f"[检查中] {msg}")


class DocumentAnnotator:
    """
    文档批注生成器
    使用 python-docx 1.2.0+ 原生批注 API，在原文段落位置插入批注。
    """

    def __init__(self):
        pass

    def generate_annotated_copy(self, original_path, format_issues, ner_data=None, ocr_data=None,
                                punctuation_errors=None, output_dir=None, suffix="_批注版"):
        """
        生成带批注的副本（增强版）

        ner_data 格式（新）:
        {
            "entities": [{"text": "...", "type": "...", "source_type": "word/ocr", "position": N}, ...],
            "summary": {...}
        }
        或兼容旧格式：
        {"ORG": [...], "PER": [...], ...}

        punctuation_errors 格式（新增）:
        list of (para_idx, char_idx, punct_char, context)
        其中 para_idx 为段落索引（0-based），char_idx 为标点在全段字符串中的索引，
        punct_char 为英文标点字符，context 为标点前后各10个字符的上下文（用于批注显示）
        """
        if not os.path.exists(original_path):
            log(f"源文件不存在: {original_path}")
            return None

        base_name = os.path.splitext(os.path.basename(original_path))[0]
        if output_dir is None:
            output_dir = os.path.dirname(os.path.abspath(original_path))

        os.makedirs(output_dir, exist_ok=True)

        output_name = f"{base_name}{suffix}.docx"
        output_path = os.path.join(output_dir, output_name)

        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except PermissionError:
                log(f"输出文件被占用，请关闭 Word 后重试: {output_path}")
                return None

        try:
            shutil.copy2(original_path, output_path)
            doc = Document(output_path)

            # 预计算表格锚点映射
            table_anchor_map = self._build_table_anchor_map(doc)

            # 找到第一个有 runs 的段落，用于全局兜底批注
            first_valid_idx = 0
            while first_valid_idx < len(doc.paragraphs) and not doc.paragraphs[first_valid_idx].runs:
                first_valid_idx += 1
            if first_valid_idx >= len(doc.paragraphs):
                first_valid_idx = 0
            log(f"全局批注锚定段落: {first_valid_idx}")

            comments_list = []

            # 1. 格式问题批注
            format_comments = self._parse_format_issues(format_issues, table_anchor_map)
            comments_list.extend(format_comments)

            # 2. NER 实体批注（增强版：按位置精确批注）
            if ner_data:
                if isinstance(ner_data, dict) and "entities" in ner_data:
                    # 新格式：带位置信息的实体列表
                    ner_comments = self._parse_ner_entities(ner_data["entities"], default_idx=first_valid_idx)
                else:
                    # 兼容旧格式
                    ner_comments = self._parse_ner_data_legacy(ner_data, default_idx=first_valid_idx)
                comments_list.extend(ner_comments)

            # 3. OCR 整体警告批注
            if ocr_data:
                ocr_comments = self._parse_ocr_data(ocr_data, default_idx=first_valid_idx)
                comments_list.extend(ocr_comments)

            # ========== 新增：标点错误精确定位批注 ==========
            if punctuation_errors:
                punct_comments = self._parse_punctuation_errors(punctuation_errors)
                comments_list.extend(punct_comments)
            # ============================================

            if not comments_list:
                log("没有需要添加的批注")
                doc.save(output_path)
                return output_path

            # 使用原生 API 添加批注
            added_comments = set()  # 避免同一位置重复添加相同内容
            for para_idx, comment_text in comments_list:
                dedup_key = f"{para_idx}|{comment_text}"
                if dedup_key in added_comments:
                    continue
                added_comments.add(dedup_key)

                if para_idx >= len(doc.paragraphs):
                    continue

                # 空段落回溯：先向前，再向后
                anchor_idx = para_idx
                while anchor_idx >= 0 and not doc.paragraphs[anchor_idx].runs:
                    anchor_idx -= 1

                if anchor_idx < 0:
                    anchor_idx = para_idx
                    while anchor_idx < len(doc.paragraphs) and not doc.paragraphs[anchor_idx].runs:
                        anchor_idx += 1
                    if anchor_idx >= len(doc.paragraphs):
                        log(f"无法为段落 {para_idx} 添加批注：文档中找不到可用锚点段落")
                        continue

                anchor_para = doc.paragraphs[anchor_idx]

                try:
                    doc.add_comment(
                        runs=anchor_para.runs,
                        text=comment_text,
                        author="暗标检查系统",
                        initials="系统"
                    )
                except Exception as e:
                    log(f"添加批注失败 (目标段落 {para_idx}, 锚点段落 {anchor_idx}): {e}")

            doc.save(output_path)
            log(f"批注副本已生成: {output_path}")
            return output_path

        except Exception as e:
            log(f"❌ 生成批注副本失败: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _build_table_anchor_map(self, doc):
        """
        构建表格锚点映射表
        """
        from docx.oxml.ns import qn

        body = doc.element.body
        table_anchors = {}

        para_idx_map = {}
        for i, para in enumerate(doc.paragraphs):
            para_idx_map[para._element] = i

        current_para_idx = -1
        table_count = 0

        for child in body:
            tag = child.tag
            if tag == qn('w:p'):
                if child in para_idx_map:
                    current_para_idx = para_idx_map[child]
            elif tag == qn('w:tbl'):
                table_count += 1
                anchor_idx = max(0, current_para_idx)
                table_anchors[table_count] = anchor_idx

        return table_anchors

    def _parse_format_issues(self, issues, table_anchor_map=None):
        """解析格式错误为批注列表"""
        if table_anchor_map is None:
            table_anchor_map = {}

        comments = []
        for issue in issues:
            # 1. 匹配 "段落 N:"
            match = re.search(r"段落\s+(\d+)", issue)
            if match:
                para_idx = int(match.group(1)) - 1
                comments.append((para_idx, issue))
                continue

            # 2. 匹配 "表格 N"
            match = re.search(r"表格\s+(\d+)", issue)
            if match:
                table_idx = int(match.group(1))
                anchor_idx = table_anchor_map.get(table_idx, 0)
                comments.append((anchor_idx, issue))
                continue

            # 3. 匹配 "第 N 节"
            match = re.search(r"第\s+(\d+)\s+节", issue)
            if match:
                comments.append((0, issue))
                continue

            # 4. 其他
            comments.append((0, issue))

        return comments

    def _parse_ner_entities(self, entities, default_idx=0):
        """
        【新增】按实体位置精确生成批注

        每个实体根据其 source_type 和 position 批注到对应段落：
        - word 类型 -> position 是段落索引，直接定位
        - ocr 类型 -> position 是图片所在段落索引，批注到图片位置
        """
        if not entities:
            return []

        # 【新增】只批注真正敏感的类型，过滤 HanLP 误识别的 OTHER/TIME/NUM 等垃圾
        VALID_TYPES = {"ORG", "PER", "LOC", "PHONE", "TEL", "EMAIL", "ID_CARD", "CREDIT_CODE", "URL"}

        # 按位置分组实体
        position_groups = {}
        for ent in entities:
            ent_type = ent.get("type", "OTHER")

            # 【关键修复】跳过无意义的类型
            if ent_type not in VALID_TYPES:
                continue

            pos = ent.get("position", default_idx)
            if pos < 0:
                pos = default_idx

            key = (ent["source_type"], pos)
            if key not in position_groups:
                position_groups[key] = []
            position_groups[key].append(ent)

        comments = []
        for (source_type, pos), group in position_groups.items():
            # 按类型分组显示
            type_groups = {}
            for ent in group:
                t = ent["type"]
                if t not in type_groups:
                    type_groups[t] = []
                type_groups[t].append(ent["text"])

            parts = []
            type_names = {
                "ORG": "机构名", "PER": "人名", "LOC": "地名",
                "PHONE": "手机号", "TEL": "座机号", "EMAIL": "邮箱",
                "ID_CARD": "身份证号", "CREDIT_CODE": "信用代码", "URL": "网址"
            }
            risk_levels = {
                "ORG": "高风险", "PER": "高风险", "LOC": "中风险",
                "PHONE": "高风险", "TEL": "高风险", "EMAIL": "中风险",
                "ID_CARD": "高风险", "CREDIT_CODE": "高风险", "URL": "中风险"
            }

            for t, texts in type_groups.items():
                unique_texts = list(dict.fromkeys(texts))[:5]  # 去重，最多显示5个
                name = type_names.get(t, t)
                risk = risk_levels.get(t, "中风险")
                parts.append(f"{name}({risk}): {', '.join(unique_texts)}")

            if parts:
                source_label = "【图片文字】" if source_type == "ocr" else "【正文文字】"
                comment_text = f"敏感信息警告{source_label} -> " + "；".join(parts)
                comments.append((pos, comment_text))

        return comments

    def _parse_ner_data_legacy(self, ner_data, default_idx=0):
        """【兼容旧格式】解析 NER 摘要字典为批注列表"""
        comments = []
        if ner_data.get("ORG"):
            orgs = ", ".join(ner_data["ORG"][:5])
            comments.append((default_idx, f"敏感信息警告：发现机构名 -> [{orgs}]，暗标文档不得出现企业标识。"))
        if ner_data.get("PER"):
            pers = ", ".join(ner_data["PER"][:5])
            comments.append((default_idx, f"敏感信息警告：发现人名 -> [{pers}]，请核实是否应删除。"))
        if ner_data.get("LOC"):
            locs = ", ".join(ner_data["LOC"][:5])
            comments.append((default_idx, f"敏感信息提示：发现地名 -> [{locs}]，请核实是否合规。"))
        return comments

    def _parse_ocr_data(self, ocr_data, default_idx=0):
        """解析 OCR 整体警告为批注列表"""
        comments = []
        if not ocr_data:
            return comments
        total_chars = sum(len(item.get("text", "")) for item in ocr_data)
        if total_chars > 50:
            comments.append((default_idx, f"图片文字警告：共 {len(ocr_data)} 张图片包含可识别文字（约 {total_chars} 字），请确认是否符合暗标要求。"))
        return comments

    # ========== 新增：解析标点错误并生成精确位置批注 ==========
    def _parse_punctuation_errors(self, punctuation_errors):
        """
        将标点错误转换为批注列表
        punctuation_errors 格式: list of (para_idx, char_idx, punct_char, context)
        批注将附加到对应段落，内容包含具体位置和上下文。
        """
        if not punctuation_errors:
            return []
        comments = []
        for para_idx, char_idx, punct_char, context in punctuation_errors:
            # 构建清晰的批注信息
            comment_text = (f"英文标点混用警告：段落第 {char_idx+1} 字符处发现英文标点 '{punct_char}'。"
                            f"附近文本：…{context}…，请使用中文标点。")
            comments.append((para_idx, comment_text))
        return comments
    # ========================================================