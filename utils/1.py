import os
from docx import Document
from docxcompose.composer import Composer
import win32com.client as win32
from pathlib import Path


def merge_word_documents_basic(folder_path, output_path):
    """
    基础方法：使用python-docx库合并（简单但可能会丢失部分格式）

    Args:
        folder_path: 包含Word文档的文件夹路径
        output_path: 输出文件路径
    """
    # 获取所有docx文件
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    docx_files.sort()  # 按文件名排序

    if not docx_files:
        print("未找到docx文件")
        return

    # 创建新文档
    merged_doc = Document()

    for file in docx_files:
        file_path = os.path.join(folder_path, file)
        print(f"正在处理: {file}")

        # 打开源文档
        source_doc = Document(file_path)

        # 复制所有段落
        for paragraph in source_doc.paragraphs:
            merged_doc.add_paragraph(paragraph.text, style=paragraph.style)

        # 分页（可选，在文档间添加分页符）
        if file != docx_files[-1]:
            merged_doc.add_page_break()

    # 保存合并后的文档
    merged_doc.save(output_path)
    print(f"合并完成！文件保存至: {output_path}")


def merge_word_documents_advanced(folder_path, output_path):
    """
    高级方法：使用win32com（保留完整格式，仅Windows）

    Args:
        folder_path: 包含Word文档的文件夹路径
        output_path: 输出文件路径
    """
    # 获取所有docx文件
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    docx_files.sort()

    if not docx_files:
        print("未找到docx文件")
        return

    # 启动Word应用程序
    word = win32.gencache.EnsureDispatch('Word.Application')
    word.Visible = False  # 不显示Word界面

    try:
        # 创建新文档
        merged_doc = word.Documents.Add()

        # 在第一页末尾添加分节符（避免格式混乱）
        merged_doc.Range().InsertBreak(7)  # 7表示分节符

        for i, file in enumerate(docx_files):
            file_path = os.path.join(folder_path, file)
            print(f"正在处理: {file}")

            # 插入文档
            if i == 0:
                # 第一个文档插入到开头
                merged_doc.Range(0, 0).InsertFile(file_path)
            else:
                # 其他文档插入到末尾
                merged_doc.Range(merged_doc.Range().End - 1, merged_doc.Range().End - 1).InsertFile(file_path)

        # 保存文档
        merged_doc.SaveAs2(output_path)
        merged_doc.Close()

        print(f"合并完成！文件保存至: {output_path}")

    except Exception as e:
        print(f"发生错误: {e}")
    finally:
        word.Quit()


def merge_word_documents_with_docxcompose(folder_path, output_path):
    """
    使用docxcompose库合并（推荐，保留格式且跨平台）

    需要安装: pip install docxcompose
    """
    from docxcompose.composer import Composer

    # 获取所有docx文件
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    docx_files.sort()

    if not docx_files:
        print("未找到docx文件")
        return

    # 打开第一个文档作为基础
    first_file = os.path.join(folder_path, docx_files[0])
    master = Document(first_file)
    composer = Composer(master)

    # 依次追加剩余文档
    for file in docx_files[1:]:
        file_path = os.path.join(folder_path, file)
        print(f"正在追加: {file}")
        doc = Document(file_path)
        composer.append(doc)

    # 保存合并后的文档
    composer.save(output_path)
    print(f"合并完成！文件保存至: {output_path}")


def merge_with_filename_prefix(folder_path, output_path, add_page_break=True):
    """
    在合并时添加文件名前缀标识，便于识别内容来源

    Args:
        folder_path: 文件夹路径
        output_path: 输出路径
        add_page_break: 是否添加分页符
    """
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    docx_files.sort()

    if not docx_files:
        print("未找到docx文件")
        return

    merged_doc = Document()

    for i, file in enumerate(docx_files):
        file_path = os.path.join(folder_path, file)
        print(f"正在处理: {file}")

        # 添加文件名作为标题（可选）
        title_paragraph = merged_doc.add_paragraph()
        title_run = title_paragraph.add_run(f"=== {file} ===")
        title_run.bold = True

        # 添加源文档内容
        source_doc = Document(file_path)
        for paragraph in source_doc.paragraphs:
            merged_doc.add_paragraph(paragraph.text, style=paragraph.style)

        # 添加分页符
        if add_page_break and i != len(docx_files) - 1:
            merged_doc.add_page_break()

    merged_doc.save(output_path)
    print(f"合并完成！文件保存至: {output_path}")


if __name__ == "__main__":
    # 设置文件夹路径和输出路径
    folder_path = r"D:\Darkbid\clients\blind_bid_tec_doc"
    output_path = r"D:\Darkbid\clients\merged_bid_document.docx"

    # 方法选择（根据需要取消注释）:

    # 方法1: 基础合并（可能丢失格式）
    # merge_word_documents_basic(folder_path, output_path)

    # 方法2: 高级合并（需要win32com，Windows专用，保留格式）
    # 注意：需要先安装 pywin32: pip install pywin32
    # merge_word_documents_advanced(folder_path, output_path)

    # 方法3: 使用docxcompose（推荐，需要安装: pip install docxcompose）
    # merge_word_documents_with_docxcompose(folder_path, output_path)

    # 方法4: 带文件名标识的合并
    merge_with_filename_prefix(folder_path, output_path, add_page_break=True)

    # 查看文件夹中的文档列表
    print("\n文件夹中的Word文档列表:")
    for f in os.listdir(folder_path):
        if f.endswith('.docx'):
            print(f"  - {f}")